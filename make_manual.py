# -*- coding: utf-8 -*-
"""生成《频谱画家使用说明》docx 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 全局默认字体
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(
    __import__("docx.oxml.ns", fromlist=["qn"]).qn("w:eastAsia"), "微软雅黑")

ACCENT = RGBColor(0xC0, 0x39, 0x2B)

def h(text, level=1):
    p = doc.add_heading(text, level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def step(no, title, body):
    p = doc.add_paragraph()
    r = p.add_run(f"第{no}步：{title}")
    r.bold = True
    doc.add_paragraph(body)

# ---------- 封面 ----------
title = doc.add_heading("无敌章鱼哥出品：频谱画家", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Invincible Squidward Presents: Spectrum Painter")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver = doc.add_paragraph("使用说明书 · v1.2.2")
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
intro = para("把文字和图片藏进声音里 —— 耳朵听不见，频谱看得见。", bold=True)
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_picture(r"e:\新建文件夹 (3)\icon\icon.png", width=Cm(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ---------- 一、软件简介 ----------
h("一、软件简介")
para("「频谱画家」是一款 Windows 桌面工具，可以将任意文字或图片编码进音频 20kHz 以上的超声波频段，"
     "与原音乐混合后导出 24-bit / 192kHz 无损文件（WAV / FLAC）。")
para("用普通耳机或音箱播放时听不出任何异常，但用 Spek 等频谱分析软件打开文件，"
     "就能看到清晰的文字或图案「悬浮」在频谱图上方。")
bullet("适用场景：音乐彩蛋、歌词可视化彩蛋、个性签名、表白密语、趣味分享")
bullet("无需安装 Python，下载 SpectroStego.exe 双击即用")

h("二、工作原理（看不懂可跳过）", 2)
para("奈奎斯特定理告诉我们：采样率的一半是可记录的最高频率。44.1kHz 的 CD 音质最多记录 22.05kHz，"
     "而人耳听力上限约 20kHz。因此将采样率提升到 192kHz 后，20–96kHz 这块频段就成了"
     "一块「听不见的空白画布」。")
para("软件把文字/图片的每一列像素映射为一帧频谱（顶部=高频，底部=低频），像素亮度决定该频点的能量，"
     "再经 IFFT（快速傅里叶逆变换）合成声波，以 -40dB 左右的低音量混入原曲，"
     "最终在频谱图上「显影」，而人耳几乎无法察觉。")

# ---------- 二、界面与操作流程 ----------
doc.add_page_break()
h("二、界面与操作流程")
para("软件界面分为左、右两部分：左侧是操作面板（按 ①②③④ 顺序操作），右侧是频谱图。")

step(1, "打开音频", "点击「① 打开音频」，支持 WAV / FLAC / OGG / MP3。"
     "建议使用无损音乐；非 192kHz 的文件会自动重采样并在界面上提示。")
step(2, "设置绘制内容", "类型选「文字」「图片 / Logo」或「二维码 / QR Code」。"
     "文字直接输入即可，字号自动适配区域高度；图片支持 PNG / JPG / BMP / WebP，"
     "透明背景自动合成；二维码自动二值化为亮底暗码，频谱截图可直接被手机扫码。")
para("重要：内容会拉伸填满你框选的区域（所见即所得）", bold=True)
bullet("框选区域的形状直接决定最终效果 —— 框什么形状，内容就是什么形状")
bullet("文字：框一个扁长区域（如 5~10 秒 × 20kHz 高）效果最佳")
bullet("图片 / Logo：按图片本身宽高比框选，深色 Logo 记得勾选「反色」")
bullet("二维码：框一个在屏幕上接近正方形的区域，扫码成功率最高")
step(3, "框选绘制区域", "在右侧频谱图上按住鼠标左键拖出一个矩形框。"
     "建议框在 20kHz 以上（例如 25–55kHz），避免与音乐本体冲突。"
     "框越高、越长，文字/图片越精细。")
step(4, "合成与验证", "点击「④ 合成到音频」，进度条走完后右侧频谱图自动刷新，"
     "可直接在软件内确认效果（与 Spek 中看到的一致），不满意可重新框选再合成。")
para("最后点击「导出 WAV」或「导出 FLAC」保存成品文件。", bold=True)

h("三、参数说明", 2)
bullet("隐藏增益（-50 ~ -20dB，默认 -40dB）：数值越大频谱图案越亮，但可能被灵敏的耳朵察觉；"
       "一般保持默认即可")
bullet("显示范围（15–96 / 0–96 / 20–60 kHz）：只影响频谱图的查看范围，不影响输出文件")
bullet("颜色（白 / 红）：仅影响预览显示，两者嵌入效果完全相同")

# ---------- 三、效果示例 ----------
doc.add_page_break()
h("三、效果示例")
para("下图左侧为文字嵌入效果（「HELLO 你好」），右侧为图片嵌入效果（章鱼哥 Logo），"
     "均位于 25–45kHz 频段：")
doc.add_picture(r"e:\新建文件夹 (3)\verify_spec_v11.png", width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("下图为二维码嵌入后从频谱中重建的效果 —— 该图可被扫码软件直接识别：")
doc.add_picture(r"e:\新建文件夹 (3)\verify_qr_square.png", width=Cm(9))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------- 四、常见问题 ----------
h("四、常见问题")
para("Q：导出后为什么听不到藏进去的内容？", bold=True)
para("A：正常现象。内容位于 20kHz 以上超声波频段，超出人耳听力范围；"
     "请用 Spek、Audition 等频谱软件查看。")
para("Q：Spek 里看不到图案？", bold=True)
para("A：① 确认 Spek 的显示上限调到 96kHz（默认可能只显示到 20kHz+）；"
     "② 确认框选区域在 20kHz 以上且合成成功；③ 尝试把隐藏增益调高到 -30dB。")
para("Q：图案模糊或变形？", bold=True)
para("A：内容会拉伸填满框选区域，图案形状由你框的区域形状决定。"
     "把框选区域拉长拉高（时间越长、频率范围越宽，可用像素越多），并按内容宽高比框选。")
para("Q：支持多长的音乐？", bold=True)
para("A：理论上不限，几分钟的歌曲均可正常处理；文件越大合成时间越长，请耐心等待进度条。")
para("Q：导出的文件很大正常吗？", bold=True)
para("A：正常。24-bit / 192kHz 是无损高规格，每分钟 WAV 约 66MB、FLAC 约为其一半。")

# ---------- 五、注意事项 ----------
h("五、注意事项")
bullet("本软件仅供学习、研究与娱乐，请勿用于侵犯版权或其他违法用途")
bullet("混音后的文件请妥善保管源文件，合成不可逆")
bullet("部分播放器/转码工具（如转 MP3）会丢弃 20kHz 以上内容，隐藏信息将丢失，分享时请直接发 WAV/FLAC 原文件")

para("")
foot = para("无敌章鱼哥出品，必属精品 🐙")
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = r"e:\新建文件夹 (3)\频谱画家使用说明.docx"
doc.save(out)
print("已生成:", out)
